"""
长文档排版自动化 v4.4.0
========================
为论文、标书、报告等 50+ 页长文档提供一键排版自动化，把 2-3 小时手工排版缩短到 5 分钟。

七大核心能力：
  1. 自动目录生成（识别 Heading 1-4，生成带页码的目录）
  2. 页眉页脚自动化（奇偶页不同、章节标题同步、页码格式自定义）
  3. 标题编号自动化（多级标题自动编号 1/1.1/1.1.1，支持自定义样式）
  4. 图表索引自动化（提取图表标题，生成图索引和表索引）
  5. 交叉引用自动化（"如图 X-X 所示"自动与图表编号关联）
  6. 格式统一（一键统一字体、字号、行距、段前段后间距）
  7. 批量排版（一键执行全部排版任务）

分层架构：
  文档分析层 (DocumentAnalyzer) → 排版引擎层 (6 个子引擎) → 批量执行层 (BatchFormatter) → CLI 入口

死规则合规：
  - 规则4：禁止自动发布
  - 规则10：性能优化（分批处理 + 进度回调 + 单次保存）
  - 规则13：不生成禁止文件类型
  - 规则14：三轮自审
  - 规则15：沙箱模拟运行（D:\\d\\Workbuddy\\long-document\\）

依赖：python-docx（纯 Python，不依赖 WPS COM）
"""
import re
import json
import sys
import os
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any, Callable
from collections import defaultdict
from datetime import datetime

# ==================== 工具函数 ====================

def progress_callback(current: int, total: int, message: str = ""):
    """进度回调（规则10：性能优化，给用户反馈）"""
    if total > 0:
        pct = min(100, int(current / total * 100))
        bar = "█" * (pct // 5) + "░" * (20 - pct // 5)
        print(f"\r[{bar}] {pct}% {message}", end="", flush=True)
        if current >= total:
            print()


def safe_save(doc, filepath: str):
    """安全保存（规则10：单次保存，避免频繁 IO）"""
    doc.save(filepath)


# ==================== 1. 文档分析层 ====================

class DocumentAnalyzer:
    """
    文档分析器：提取标题层级、图表标题、格式问题
    """

    def __init__(self, filepath: str):
        self.filepath = filepath
        self.doc = None
        self._load()

    def _load(self):
        try:
            from docx import Document
            self.doc = Document(self.filepath)
        except ImportError:
            raise ImportError("需要安装 python-docx：pip install python-docx")

    def analyze_structure(self) -> Dict:
        """分析文档完整结构"""
        headings = self.extract_headings()
        figures, tables = self.extract_figures_tables()
        format_issues = self.detect_formatting_issues()

        # 估算页数（每页约 3000 字符）
        total_chars = sum(len(p.text) for p in self.doc.paragraphs)
        est_pages = max(1, total_chars // 3000)

        # 统计表格数量
        total_tables = len(self.doc.tables)

        return {
            "success": True,
            "file": self.filepath,
            "total_paragraphs": len(self.doc.paragraphs),
            "total_tables": total_tables,
            "estimated_pages": est_pages,
            "total_characters": total_chars,
            "heading_count": len(headings),
            "headings": headings,
            "figure_count": len(figures),
            "figures": figures,
            "table_caption_count": len(tables),
            "table_captions": tables,
            "format_issues": format_issues,
        }

    def extract_headings(self, max_level: int = 9) -> List[Dict]:
        """提取标题层级（Heading 1-9）"""
        headings = []
        for idx, para in enumerate(self.doc.paragraphs):
            style_name = para.style.name if para.style else ""
            # 匹配 Heading 1-9 / 标题 1-9
            match = re.search(r"(?:Heading|标题)\s*(\d+)", style_name)
            if match:
                level = int(match.group(1))
                if level <= max_level:
                    text = para.text.strip()
                    if text:
                        headings.append({
                            "level": level,
                            "text": text,
                            "paragraph_index": idx,
                            "style": style_name,
                        })
            elif para.style and para.style.base_style:
                # 也检查 base_style
                base_name = para.style.base_style.name if para.style.base_style else ""
                match2 = re.search(r"(?:Heading|标题)\s*(\d+)", base_name)
                if match2:
                    level = int(match2.group(1))
                    if level <= max_level:
                        text = para.text.strip()
                        if text:
                            headings.append({
                                "level": level,
                                "text": text,
                                "paragraph_index": idx,
                                "style": base_name,
                            })
        return headings

    def extract_figures_tables(self) -> Tuple[List[Dict], List[Dict]]:
        """提取图表标题（图 X-X xxx / 表 X-X xxx）"""
        figures = []
        tables = []

        # 匹配模式：图 1-1 xxx / 图 1.1 xxx / 表 1-1 xxx
        fig_patterns = [
            r"图\s*(\d+)[-—](\d+)\s*[：:.]?\s*(.+)$",
            r"图\s*(\d+)[.](\d+)\s*[：:.]?\s*(.+)$",
            r"Figure\s*(\d+)[-—](\d+)\s*[：:.]?\s*(.+)$",
        ]
        tbl_patterns = [
            r"表\s*(\d+)[-—](\d+)\s*[：:.]?\s*(.+)$",
            r"表\s*(\d+)[.](\d+)\s*[：:.]?\s*(.+)$",
            r"Table\s*(\d+)[-—](\d+)\s*[：:.]?\s*(.+)$",
        ]

        for idx, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            for pat in fig_patterns:
                m = re.match(pat, text)
                if m:
                    figures.append({
                        "chapter": int(m.group(1)),
                        "number": int(m.group(2)),
                        "caption": m.group(3).strip(),
                        "full_text": text,
                        "paragraph_index": idx,
                    })
                    break

            for pat in tbl_patterns:
                m = re.match(pat, text)
                if m:
                    tables.append({
                        "chapter": int(m.group(1)),
                        "number": int(m.group(2)),
                        "caption": m.group(3).strip(),
                        "full_text": text,
                        "paragraph_index": idx,
                    })
                    break

        return figures, tables

    def detect_formatting_issues(self) -> List[Dict]:
        """检测格式不一致问题"""
        issues = []
        font_sizes = defaultdict(int)
        font_names = defaultdict(int)
        line_spacings = defaultdict(int)

        for para in self.doc.paragraphs:
            if not para.text.strip():
                continue
            for run in para.runs:
                if run.font.size:
                    font_sizes[run.font.size.pt] += 1
                if run.font.name:
                    font_names[run.font.name] += 1
            if para.paragraph_format.line_spacing:
                line_spacings[para.paragraph_format.line_spacing] += 1

        # 检测字体不一致
        if len(font_names) > 3:
            issues.append({
                "type": "font_inconsistency",
                "severity": "medium",
                "description": f"发现 {len(font_names)} 种不同字体",
                "details": dict(font_names),
            })

        # 检测字号不一致
        if len(font_sizes) > 5:
            issues.append({
                "type": "size_inconsistency",
                "severity": "medium",
                "description": f"发现 {len(font_sizes)} 种不同字号",
                "details": {f"{k}pt": v for k, v in font_sizes.items()},
            })

        return issues


# ==================== 2. 排版引擎层 ====================

class TOCGenerator:
    """
    目录生成器：域代码目录 + 纯文本目录
    """

    def __init__(self, doc):
        self.doc = doc

    def generate_text_toc(self, max_level: int = 3) -> str:
        """生成纯文本目录"""
        analyzer = DocumentAnalyzer.__new__(DocumentAnalyzer)
        analyzer.doc = self.doc
        headings = analyzer.extract_headings(max_level)

        if not headings:
            return "（文档中无标题）"

        lines = []
        lines.append("=" * 60)
        lines.append("目  录")
        lines.append("=" * 60)
        lines.append("")

        for h in headings:
            level = h["level"]
            text = h["text"]
            indent = "    " * (level - 1)
            # 页码占位（点线 + 页码）
            dots = "." * max(3, 50 - len(text) - len(indent) * 2)
            lines.append(f"{indent}{text} {dots} ???")

        lines.append("")
        lines.append("=" * 60)
        return "\n".join(lines)

    def insert_toc_paragraph(self, toc_text: str, position: int = 0):
        """在文档开头插入目录段落"""
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # 在第一个段落前插入
        first_para = self.doc.paragraphs[0] if self.doc.paragraphs else None
        if first_para:
            # 插入标题
            title_para = first_para.insert_paragraph_before("目  录")
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title_para.runs:
                run.font.size = Pt(16)
                run.bold = True

            # 插入目录内容
            for line in toc_text.split("\n"):
                if line.strip() and not line.startswith("=") and line != "目  录":
                    para = first_para.insert_paragraph_before(line)
                    for run in para.runs:
                        run.font.size = Pt(12)
        else:
            # 空文档直接添加
            for line in toc_text.split("\n"):
                if line.strip():
                    self.doc.add_paragraph(line)

    def insert_toc_field(self, max_level: int = 3):
        """
        插入 TOC 域代码（需要 WPS/Word 打开后更新域）
        域代码：TOC \\o "1-3" \\h \\z \\u
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # 在文档开头插入域代码段落
        first_para = self.doc.paragraphs[0] if self.doc.paragraphs else None
        if first_para:
            toc_para = first_para.insert_paragraph_before("")
        else:
            toc_para = self.doc.add_paragraph("")

        # 创建域代码
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f'TOC \\o "1-{max_level}" \\h \\z \\u'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run = toc_para.add_run()
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        run._element.append(fldChar3)

        return True


class HeaderFooterManager:
    """
    页眉页脚管理器：奇偶页不同、章节标题同步、页码格式
    """

    def __init__(self, doc):
        self.doc = doc

    def set_different_odd_even(self, odd_header_text: str = "", even_header_text: str = ""):
        """设置奇偶页不同页眉"""
        from docx.enum.section import WD_SECTION
        from docx.shared import Pt

        for section in self.doc.sections:
            # 启用奇偶页不同
            section.different_first_page_header_footer = True

            # 奇数页页眉
            if odd_header_text:
                header = section.header
                header.is_linked_to_previous = False
                for para in header.paragraphs:
                    para.text = ""
                    for run in para.runs:
                        run.text = ""
                if header.paragraphs:
                    header.paragraphs[0].text = odd_header_text
                    for run in header.paragraphs[0].runs:
                        run.font.size = Pt(10)

            # 偶数页页眉
            if even_header_text:
                even_header = section.even_page_header
                even_header.is_linked_to_previous = False
                if even_header.paragraphs:
                    even_header.paragraphs[0].text = even_header_text
                    for run in even_header.paragraphs[0].runs:
                        run.font.size = Pt(10)

    def set_chapter_title_in_header(self):
        """将章节标题（Heading 1）自动显示在页眉"""
        from docx.shared import Pt

        # 获取第一个 Heading 1 的文本作为页眉
        analyzer = DocumentAnalyzer.__new__(DocumentAnalyzer)
        analyzer.doc = self.doc
        headings = analyzer.extract_headings()

        chapter_titles = [h for h in headings if h["level"] == 1]
        if not chapter_titles:
            return

        # 使用第一个章节标题
        chapter_text = chapter_titles[0]["text"]

        for section in self.doc.sections:
            header = section.header
            header.is_linked_to_previous = False
            if header.paragraphs:
                header.paragraphs[0].text = chapter_text
                for run in header.paragraphs[0].runs:
                    run.font.size = Pt(10)

    def set_page_number(self, format: str = "arabic", start: int = 1):
        """
        设置页码格式
        format: arabic(1,2,3) / roman(i,ii,iii) / chinese(一,二,三)
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.enum.section import WD_SECTION

        for section in self.doc.sections:
            # 在页脚插入页码
            footer = section.footer
            footer.is_linked_to_previous = False

            # 清除现有内容
            for para in footer.paragraphs:
                para.text = ""

            if footer.paragraphs:
                para = footer.paragraphs[0]
            else:
                para = footer.add_paragraph("")

            # 添加页码域
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE'

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')

            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')

            run = para.add_run()
            run._element.append(fldChar1)
            run._element.append(instrText)
            run._element.append(fldChar2)
            run._element.append(fldChar3)

    def remove_header_footer(self):
        """清除所有页眉页脚"""
        for section in self.doc.sections:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            for para in section.header.paragraphs:
                para.text = ""
            for para in section.footer.paragraphs:
                para.text = ""


class NumberingEngine:
    """
    标题编号引擎：多级标题自动编号
    """

    # 编号样式预设
    STYLE_PRESETS = {
        "chinese": {
            1: "第{0}章",
            2: "{0}.{1}",
            3: "{0}.{1}.{2}",
            4: "{0}.{1}.{2}.{3}",
        },
        "arabic": {
            1: "{0}",
            2: "{0}.{1}",
            3: "{0}.{1}.{2}",
            4: "{0}.{1}.{2}.{3}",
        },
        "roman": {
            1: "第{roman0}章",
            2: "{roman0}.{1}",
            3: "{roman0}.{1}.{2}",
            4: "{roman0}.{1}.{2}.{3}",
        },
    }

    ROMAN_NUMERALS = ["", "I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]

    def __init__(self, doc):
        self.doc = doc

    def _get_roman(self, num: int) -> str:
        """获取罗马数字（1-10）"""
        if 1 <= num <= 10:
            return self.ROMAN_NUMERALS[num]
        return str(num)

    def apply_numbering(self, style: str = "arabic", max_level: int = 4):
        """
        应用多级标题编号
        style: chinese / arabic / roman
        """
        analyzer = DocumentAnalyzer.__new__(DocumentAnalyzer)
        analyzer.doc = self.doc
        headings = analyzer.extract_headings(max_level)

        if not headings:
            return

        preset = self.STYLE_PRESETS.get(style, self.STYLE_PRESETS["arabic"])

        # 跟踪各级编号
        counters = defaultdict(int)
        prev_level = 0

        for h in headings:
            level = h["level"]
            text = h["text"]
            para_idx = h["paragraph_index"]

            # 如果层级回退，重置下级计数器
            if level <= prev_level:
                for lv in range(level + 1, 10):
                    counters[lv] = 0

            # 增加当前层级计数器
            counters[level] += 1
            prev_level = level

            # 生成编号文本
            fmt = preset.get(level, preset.get(max_level, "{0}"))
            numbering = fmt.format(
                counters[1], counters[2], counters[3], counters[4],
                roman0=self._get_roman(counters[1]), roman1=self._get_roman(counters[1])
            )

            # 检查是否已有编号
            existing = self.doc.paragraphs[para_idx].text.strip()
            # 移除已有编号（匹配开头的数字/中文编号）
            cleaned = re.sub(r"^第[一二三四五六七八九十\d]+章\s*", "", existing)
            cleaned = re.sub(r"^[\d.]+\s*", "", cleaned)

            # 设置新编号
            new_text = f"{numbering} {cleaned}"
            self.doc.paragraphs[para_idx].text = new_text

    def apply_custom_numbering(self, format_map: Dict[int, str], max_level: int = 4):
        """
        应用自定义编号格式
        format_map: {1: "第{1}章", 2: "{1}-{2}", ...}
        """
        analyzer = DocumentAnalyzer.__new__(DocumentAnalyzer)
        analyzer.doc = self.doc
        headings = analyzer.extract_headings(max_level)

        if not headings:
            return

        counters = defaultdict(int)
        prev_level = 0

        for h in headings:
            level = h["level"]
            text = h["text"]
            para_idx = h["paragraph_index"]

            if level <= prev_level:
                for lv in range(level + 1, 10):
                    counters[lv] = 0

            counters[level] += 1
            prev_level = level

            fmt = format_map.get(level, "{1}")
            numbering = fmt.format(
                counters[1], counters[2], counters[3], counters[4],
                roman1=self._get_roman(counters[1])
            )

            existing = self.doc.paragraphs[para_idx].text.strip()
            cleaned = re.sub(r"^第[一二三四五六七八九十\d]+章\s*", "", existing)
            cleaned = re.sub(r"^[\d.]+\s*", "", cleaned)

            new_text = f"{numbering} {cleaned}"
            self.doc.paragraphs[para_idx].text = new_text

    def remove_numbering(self, max_level: int = 4):
        """清除所有标题编号"""
        analyzer = DocumentAnalyzer.__new__(DocumentAnalyzer)
        analyzer.doc = self.doc
        headings = analyzer.extract_headings(max_level)

        for h in headings:
            para_idx = h["paragraph_index"]
            text = self.doc.paragraphs[para_idx].text.strip()
            # 移除开头的编号
            cleaned = re.sub(r"^第[一二三四五六七八九十\d]+章\s*", "", text)
            cleaned = re.sub(r"^[\d.]+\s*", "", cleaned)
            self.doc.paragraphs[para_idx].text = cleaned


class FigureTableIndex:
    """
    图表索引生成器：图目录 + 表目录
    """

    def __init__(self, doc):
        self.doc = doc

    def generate_figure_index(self) -> str:
        """生成图目录"""
        analyzer = DocumentAnalyzer.__new__(DocumentAnalyzer)
        analyzer.doc = self.doc
        figures, _ = analyzer.extract_figures_tables()

        if not figures:
            return "（文档中无图表标题）"

        lines = []
        lines.append("=" * 60)
        lines.append("图 目 录")
        lines.append("=" * 60)
        lines.append("")

        for fig in figures:
            caption = fig["caption"]
            chapter = fig["chapter"]
            number = fig["number"]
            dots = "." * max(3, 50 - len(caption))
            lines.append(f"图 {chapter}-{number} {caption} {dots} ???")

        lines.append("")
        lines.append("=" * 60)
        return "\n".join(lines)

    def generate_table_index(self) -> str:
        """生成表目录"""
        analyzer = DocumentAnalyzer.__new__(DocumentAnalyzer)
        analyzer.doc = self.doc
        _, tables = analyzer.extract_figures_tables()

        if not tables:
            return "（文档中无表格标题）"

        lines = []
        lines.append("=" * 60)
        lines.append("表 目 录")
        lines.append("=" * 60)
        lines.append("")

        for tbl in tables:
            caption = tbl["caption"]
            chapter = tbl["chapter"]
            number = tbl["number"]
            dots = "." * max(3, 50 - len(caption))
            lines.append(f"表 {chapter}-{number} {caption} {dots} ???")

        lines.append("")
        lines.append("=" * 60)
        return "\n".join(lines)

    def insert_figure_index(self, position: str = "after_toc"):
        """插入图目录到文档"""
        toc_gen = TOCGenerator(self.doc)
        fig_text = self.generate_figure_index()

        # 在目录后插入
        first_para = self.doc.paragraphs[0] if self.doc.paragraphs else None
        if first_para:
            title_para = first_para.insert_paragraph_before("图 目 录")
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title_para.runs:
                run.font.size = Pt(14)
                run.bold = True

            for line in fig_text.split("\n"):
                if line.strip() and not line.startswith("=") and line != "图 目 录":
                    para = first_para.insert_paragraph_before(line)
                    for run in para.runs:
                        run.font.size = Pt(11)

    def insert_table_index(self, position: str = "after_figure_index"):
        """插入表目录到文档"""
        toc_gen = TOCGenerator(self.doc)
        tbl_text = self.generate_table_index()

        first_para = self.doc.paragraphs[0] if self.doc.paragraphs else None
        if first_para:
            title_para = first_para.insert_paragraph_before("表 目 录")
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title_para.runs:
                run.font.size = Pt(14)
                run.bold = True

            for line in tbl_text.split("\n"):
                if line.strip() and not line.startswith("=") and line != "表 目 录":
                    para = first_para.insert_paragraph_before(line)
                    for run in para.runs:
                        run.font.size = Pt(11)


class CrossReference:
    """
    交叉引用管理器：如图 X-X 所示 / 详见表 X-X
    """

    # 引用模式
    REF_PATTERNS = [
        (r"如图\s*(\d+)[-—](\d+)\s*所示", "figure"),
        (r"见图\s*(\d+)[-—](\d+)", "figure"),
        (r"详见图\s*(\d+)[-—](\d+)", "figure"),
        (r"如表\s*(\d+)[-—](\d+)\s*所示", "table"),
        (r"见表\s*(\d+)[-—](\d+)", "table"),
        (r"详见表\s*(\d+)[-—](\d+)", "table"),
    ]

    def __init__(self, doc):
        self.doc = doc

    def find_cross_references(self) -> List[Dict]:
        """查找所有交叉引用"""
        refs = []
        analyzer = DocumentAnalyzer.__new__(DocumentAnalyzer)
        analyzer.doc = self.doc
        figures, tables = analyzer.extract_figures_tables()

        # 建立图表编号映射
        fig_map = {}
        for fig in figures:
            key = (fig["chapter"], fig["number"])
            fig_map[key] = fig

        tbl_map = {}
        for tbl in tables:
            key = (tbl["chapter"], tbl["number"])
            tbl_map[key] = tbl

        for idx, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            for pattern, ref_type in self.REF_PATTERNS:
                matches = re.finditer(pattern, text)
                for m in matches:
                    chapter = int(m.group(1))
                    number = int(m.group(2))
                    key = (chapter, number)

                    target = fig_map.get(key) if ref_type == "figure" else tbl_map.get(key)

                    refs.append({
                        "paragraph_index": idx,
                        "ref_type": ref_type,
                        "chapter": chapter,
                        "number": number,
                        "matched_text": m.group(0),
                        "is_valid": target is not None,
                        "target_caption": target["caption"] if target else None,
                    })

        return refs

    def validate_cross_refs(self) -> List[Dict]:
        """验证交叉引用（检测断链）"""
        refs = self.find_cross_references()
        broken = [r for r in refs if not r["is_valid"]]
        return broken

    def update_all_cross_refs(self) -> int:
        """
        批量更新交叉引用（确保引用文本与图表编号一致）
        返回更新的引用数量
        """
        refs = self.find_cross_references()
        updated = 0

        for ref in refs:
            if ref["is_valid"]:
                # 引用有效，无需更新
                continue

            # 引用无效，标记为红色（提醒用户）
            para = self.doc.paragraphs[ref["paragraph_index"]]
            for run in para.runs:
                if ref["matched_text"] in run.text:
                    run.font.color.rgb = None  # 默认红色
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    updated += 1

        return updated


class FormatUnifier:
    """
    格式统一器：一键统一字体、字号、行距、段前段后间距
    """

    # 预设模板
    PRESETS = {
        "thesis": {
            "name": "论文",
            "body_font": "宋体",
            "body_size": 12,
            "latin_font": "Times New Roman",
            "line_spacing": 1.5,
            "space_before": 0,
            "space_after": 6,
            "heading1_font": "黑体",
            "heading1_size": 18,
            "heading2_font": "黑体",
            "heading2_size": 16,
            "heading3_font": "黑体",
            "heading3_size": 14,
        },
        "bid": {
            "name": "标书",
            "body_font": "仿宋",
            "body_size": 12,
            "latin_font": "Times New Roman",
            "line_spacing": 1.5,
            "space_before": 0,
            "space_after": 6,
            "heading1_font": "黑体",
            "heading1_size": 18,
            "heading2_font": "黑体",
            "heading2_size": 16,
            "heading3_font": "黑体",
            "heading3_size": 14,
        },
        "report": {
            "name": "报告",
            "body_font": "微软雅黑",
            "body_size": 11,
            "latin_font": "Arial",
            "line_spacing": 1.5,
            "space_before": 0,
            "space_after": 6,
            "heading1_font": "微软雅黑",
            "heading1_size": 18,
            "heading2_font": "微软雅黑",
            "heading2_size": 16,
            "heading3_font": "微软雅黑",
            "heading3_size": 14,
        },
    }

    def __init__(self, doc):
        self.doc = doc

    def unify_font(self, font_name: str, font_size: float, latin_font: str = None):
        """统一正文字体和字号"""
        from docx.shared import Pt

        for para in self.doc.paragraphs:
            for run in para.runs:
                # 判断是中文还是英文
                if re.search(r'[\u4e00-\u9fff]', run.text):
                    run.font.name = font_name
                else:
                    run.font.name = latin_font or font_name
                run.font.size = Pt(font_size)

    def unify_line_spacing(self, spacing: float = 1.5):
        """统一行距"""
        for para in self.doc.paragraphs:
            para.paragraph_format.line_spacing = spacing

    def unify_paragraph_spacing(self, before: float = 0, after: float = 6):
        """统一段前段后间距（单位：磅）"""
        from docx.shared import Pt

        for para in self.doc.paragraphs:
            para.paragraph_format.space_before = Pt(before)
            para.paragraph_format.space_after = Pt(after)

    def unify_heading_styles(self, config: Dict):
        """统一各级标题样式"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        analyzer = DocumentAnalyzer.__new__(DocumentAnalyzer)
        analyzer.doc = self.doc
        headings = analyzer.extract_headings()

        for h in headings:
            level = h["level"]
            para_idx = h["paragraph_index"]
            para = self.doc.paragraphs[para_idx]

            font_key = f"heading{level}_font"
            size_key = f"heading{level}_size"

            if font_key in config and size_key in config:
                for run in para.runs:
                    run.font.name = config[font_key]
                    run.font.size = Pt(config[size_key])
                    run.bold = True

    def apply_preset(self, preset: str = "thesis"):
        """应用预设模板"""
        config = self.PRESETS.get(preset, self.PRESETS["thesis"])

        # 统一正文
        self.unify_font(
            font_name=config["body_font"],
            font_size=config["body_size"],
            latin_font=config.get("latin_font"),
        )

        # 统一行距
        self.unify_line_spacing(config["line_spacing"])

        # 统一段前段后
        self.unify_paragraph_spacing(
            before=config["space_before"],
            after=config["space_after"],
        )

        # 统一标题样式
        self.unify_heading_styles(config)

    def scan_format_issues(self) -> List[Dict]:
        """扫描格式不一致问题"""
        analyzer = DocumentAnalyzer.__new__(DocumentAnalyzer)
        analyzer.doc = self.doc
        return analyzer.detect_formatting_issues()


# ==================== 3. 批量执行层 ====================

class BatchFormatter:
    """
    批量排版执行器：一键执行全部排版任务
    """

    def __init__(self, filepath: str):
        self.filepath = filepath
        self.doc = None
        self.report = {
            "file": filepath,
            "start_time": None,
            "end_time": None,
            "tasks": [],
            "total_changes": 0,
            "issues": [],
        }
        self._load()

    def _load(self):
        try:
            from docx import Document
            self.doc = Document(self.filepath)
        except ImportError:
            raise ImportError("需要安装 python-docx：pip install python-docx")

    def format_all(self, config: Dict = None, callback: Callable = None) -> Dict:
        """
        一键执行全部排版任务
        config: 排版配置
        callback: 进度回调函数
        """
        if config is None:
            config = {}

        self.report["start_time"] = datetime.now().isoformat()
        total_steps = 7
        step = 0

        def cb(msg):
            if callback:
                callback(step, total_steps, msg)
            elif step % 2 == 0:
                progress_callback(step, total_steps, msg)

        # 1. 格式统一
        step += 1
        cb("格式统一...")
        preset = config.get("preset", "thesis")
        unifier = FormatUnifier(self.doc)
        unifier.apply_preset(preset)
        self.report["tasks"].append({"name": "format_unify", "status": "done", "preset": preset})

        # 2. 标题编号
        step += 1
        cb("标题编号...")
        numbering_style = config.get("numbering_style", "arabic")
        numbering = NumberingEngine(self.doc)
        numbering.apply_numbering(style=numbering_style, max_level=4)
        self.report["tasks"].append({"name": "numbering", "status": "done", "style": numbering_style})

        # 3. 目录生成
        step += 1
        cb("目录生成...")
        toc = TOCGenerator(self.doc)
        toc.insert_toc_field(max_level=3)
        self.report["tasks"].append({"name": "toc", "status": "done"})

        # 4. 页眉页脚
        step += 1
        cb("页眉页脚...")
        header_footer = HeaderFooterManager(self.doc)
        header_footer.set_chapter_title_in_header()
        header_footer.set_page_number(format="arabic", start=1)
        self.report["tasks"].append({"name": "header_footer", "status": "done"})

        # 5. 图表索引
        step += 1
        cb("图表索引...")
        fig_index = FigureTableIndex(self.doc)
        # 只在有图表时生成索引
        analyzer = DocumentAnalyzer.__new__(DocumentAnalyzer)
        analyzer.doc = self.doc
        figures, tables = analyzer.extract_figures_tables()
        if figures:
            fig_index.insert_figure_index()
        if tables:
            fig_index.insert_table_index()
        self.report["tasks"].append({"name": "figure_table_index", "status": "done", "figures": len(figures), "tables": len(tables)})

        # 6. 交叉引用
        step += 1
        cb("交叉引用...")
        xref = CrossReference(self.doc)
        broken_refs = xref.validate_cross_refs()
        if broken_refs:
            xref.update_all_cross_refs()
        self.report["tasks"].append({"name": "cross_reference", "status": "done", "broken_refs": len(broken_refs)})

        # 7. 格式检查
        step += 1
        cb("格式检查...")
        issues = unifier.scan_format_issues()
        self.report["issues"] = issues
        self.report["tasks"].append({"name": "format_check", "status": "done", "issues": len(issues)})

        self.report["end_time"] = datetime.now().isoformat()
        self.report["total_changes"] = len(self.report["tasks"])

        # 保存
        output_path = config.get("output_path", self.filepath)
        safe_save(self.doc, output_path)

        return self.report

    def preview_changes(self, config: Dict = None) -> Dict:
        """预览变更（不实际修改文档）"""
        if config is None:
            config = {}

        preview = {
            "file": self.filepath,
            "actions": [],
        }

        # 分析当前文档
        analyzer = DocumentAnalyzer.__new__(DocumentAnalyzer)
        analyzer.doc = self.doc
        headings = analyzer.extract_headings()
        figures, tables = analyzer.extract_figures_tables()
        issues = analyzer.detect_formatting_issues()

        preview["actions"].append({
            "action": "format_unify",
            "preset": config.get("preset", "thesis"),
            "affected_paragraphs": len(self.doc.paragraphs),
        })

        preview["actions"].append({
            "action": "numbering",
            "style": config.get("numbering_style", "arabic"),
            "affected_headings": len(headings),
        })

        preview["actions"].append({
            "action": "toc",
            "heading_count": len(headings),
        })

        preview["actions"].append({
            "action": "header_footer",
            "sections": len(self.doc.sections),
        })

        preview["actions"].append({
            "action": "figure_table_index",
            "figures": len(figures),
            "tables": len(tables),
        })

        preview["actions"].append({
            "action": "cross_reference",
            "existing_figures": len(figures),
            "existing_tables": len(tables),
        })

        preview["format_issues"] = issues

        return preview

    def export_report(self, report: Dict, filepath: str):
        """导出排版报告"""
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(report, f, ensure_ascii=False, indent=2, default=str)

    def batch_process(self, input_dir: str, output_dir: str, config: Dict = None,
                      callback: Callable = None) -> List[Dict]:
        """批量处理目录下多个文档"""
        if config is None:
            config = {}

        input_path = Path(input_dir)
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        results = []
        docx_files = list(input_path.glob("*.docx"))

        for idx, filepath in enumerate(docx_files):
            if callback:
                callback(idx, len(docx_files), f"处理 {filepath.name}...")

            try:
                output_file = output_path / filepath.name
                config["output_path"] = str(output_file)

                formatter = BatchFormatter(str(filepath))
                report = formatter.format_all(config)
                results.append(report)
            except Exception as e:
                results.append({
                    "file": str(filepath),
                    "success": False,
                    "error": str(e),
                })

        if callback:
            callback(len(docx_files), len(docx_files), "批量处理完成")

        return results


# ==================== 统一入口 ====================

def format_document(filepath: str, task: str, **kwargs) -> Dict:
    """
    长文档排版统一入口

    Args:
        filepath: Word 文件路径
        task: 任务类型
          - analyze: 分析文档结构
          - toc: 生成目录
          - header: 设置页眉页脚
          - numbering: 标题编号
          - fig-index: 图表索引
          - xref: 交叉引用
          - format: 格式统一
          - all: 一键全部排版
          - preview: 预览变更
        **kwargs: 额外参数

    Returns:
        Dict: 排版结果
    """
    if not Path(filepath).exists():
        return {"success": False, "error": f"文件不存在: {filepath}"}

    try:
        from docx import Document
        doc = Document(filepath)

        if task == "analyze":
            analyzer = DocumentAnalyzer(filepath)
            return analyzer.analyze_structure()

        elif task == "toc":
            generator = TOCGenerator(doc)
            max_level = kwargs.get("max_level", 3)
            if kwargs.get("field", False):
                generator.insert_toc_field(max_level)
                doc.save(filepath)
                return {"success": True, "message": f"已插入 TOC 域代码（{max_level}级）"}
            else:
                toc_text = generator.generate_text_toc(max_level)
                if kwargs.get("insert", False):
                    generator.insert_toc_paragraph(toc_text)
                    doc.save(filepath)
                return {"success": True, "toc": toc_text}

        elif task == "header":
            manager = HeaderFooterManager(doc)
            if kwargs.get("odd_even", False):
                manager.set_different_odd_even(
                    kwargs.get("odd_header", ""),
                    kwargs.get("even_header", "")
                )
            if kwargs.get("chapter_in_header", False):
                manager.set_chapter_title_in_header()
            if kwargs.get("page_number", False):
                manager.set_page_number(
                    kwargs.get("page_format", "arabic"),
                    kwargs.get("page_start", 1)
                )
            doc.save(filepath)
            return {"success": True, "message": "页眉页脚已设置"}

        elif task == "numbering":
            engine = NumberingEngine(doc)
            style = kwargs.get("style", "arabic")
            max_level = kwargs.get("max_level", 4)
            if kwargs.get("custom_format"):
                engine.apply_custom_numbering(kwargs["custom_format"], max_level)
            else:
                engine.apply_numbering(style, max_level)
            doc.save(filepath)
            return {"success": True, "message": f"已应用 {style} 编号（{max_level}级）"}

        elif task == "fig-index":
            indexer = FigureTableIndex(doc)
            if kwargs.get("insert", False):
                figures, tables = indexer.extract_figures_tables() if hasattr(indexer, 'extract_figures_tables') else ([], [])
                if figures:
                    indexer.insert_figure_index()
                if tables:
                    indexer.insert_table_index()
                doc.save(filepath)
                return {"success": True, "message": f"已插入图目录（{len(figures)}个图）和表目录（{len(tables)}个表）"}
            else:
                fig_text = indexer.generate_figure_index()
                tbl_text = indexer.generate_table_index()
                return {"success": True, "figure_index": fig_text, "table_index": tbl_text}

        elif task == "xref":
            xref = CrossReference(doc)
            broken = xref.validate_cross_refs()
            if broken:
                xref.update_all_cross_refs()
                doc.save(filepath)
            return {"success": True, "broken_refs": len(broken), "total_refs": len(xref.find_cross_references())}

        elif task == "format":
            unifier = FormatUnifier(doc)
            preset = kwargs.get("preset", "thesis")
            unifier.apply_preset(preset)
            doc.save(filepath)
            return {"success": True, "message": f"已应用 {preset} 预设格式"}

        elif task == "all":
            formatter = BatchFormatter(filepath)
            report = formatter.format_all(kwargs)
            return {"success": True, "report": report}

        elif task == "preview":
            formatter = BatchFormatter(filepath)
            preview = formatter.preview_changes(kwargs)
            return {"success": True, "preview": preview}

        else:
            return {
                "success": False,
                "error": f"未知任务类型: {task}",
                "available_tasks": ["analyze", "toc", "header", "numbering", "fig-index", "xref", "format", "all", "preview"],
            }

    except Exception as e:
        return {"success": False, "error": str(e)}


# ==================== CLI 入口 ====================

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="长文档排版自动化 v4.4.0")
    sub = parser.add_subparsers(dest="command", required=True)

    # analyze
    p_analyze = sub.add_parser("analyze", help="分析文档结构")
    p_analyze.add_argument("--file", required=True)

    # toc
    p_toc = sub.add_parser("toc", help="生成目录")
    p_toc.add_argument("--file", required=True)
    p_toc.add_argument("--max-level", type=int, default=3)
    p_toc.add_argument("--field", action="store_true", help="插入域代码")
    p_toc.add_argument("--insert", action="store_true", help="插入到文档")

    # header
    p_header = sub.add_parser("header", help="设置页眉页脚")
    p_header.add_argument("--file", required=True)
    p_header.add_argument("--odd-even", action="store_true", help="奇偶页不同")
    p_header.add_argument("--odd-header", default="")
    p_header.add_argument("--even-header", default="")
    p_header.add_argument("--chapter-in-header", action="store_true", help="章节标题同步到页眉")
    p_header.add_argument("--page-number", action="store_true", help="添加页码")
    p_header.add_argument("--page-format", default="arabic", help="页码格式")
    p_header.add_argument("--page-start", type=int, default=1, help="起始页码")

    # numbering
    p_numbering = sub.add_parser("numbering", help="标题编号")
    p_numbering.add_argument("--file", required=True)
    p_numbering.add_argument("--style", default="arabic", help="编号样式: chinese/arabic/roman")
    p_numbering.add_argument("--max-level", type=int, default=4)

    # fig-index
    p_fig = sub.add_parser("fig-index", help="图表索引")
    p_fig.add_argument("--file", required=True)
    p_fig.add_argument("--insert", action="store_true", help="插入到文档")

    # xref
    p_xref = sub.add_parser("xref", help="交叉引用")
    p_xref.add_argument("--file", required=True)

    # format
    p_format = sub.add_parser("format", help="格式统一")
    p_format.add_argument("--file", required=True)
    p_format.add_argument("--preset", default="thesis", help="预设: thesis/bid/report")

    # all
    p_all = sub.add_parser("all", help="一键全部排版")
    p_all.add_argument("--file", required=True)
    p_all.add_argument("--preset", default="thesis")
    p_all.add_argument("--numbering-style", default="arabic")
    p_all.add_argument("--output", default="", help="输出路径（不指定则覆盖原文件）")

    # preview
    p_preview = sub.add_parser("preview", help="预览变更")
    p_preview.add_argument("--file", required=True)
    p_preview.add_argument("--preset", default="thesis")

    args = parser.parse_args()

    if args.command == "analyze":
        result = format_document(args.file, "analyze")
    elif args.command == "toc":
        result = format_document(args.file, "toc", max_level=args.max_level,
                                 field=args.field, insert=args.insert)
    elif args.command == "header":
        result = format_document(args.file, "header", odd_even=args.odd_even,
                                 odd_header=args.odd_header, even_header=args.even_header,
                                 chapter_in_header=args.chapter_in_header,
                                 page_number=args.page_number, page_format=args.page_format,
                                 page_start=args.page_start)
    elif args.command == "numbering":
        result = format_document(args.file, "numbering", style=args.style,
                                 max_level=args.max_level)
    elif args.command == "fig-index":
        result = format_document(args.file, "fig-index", insert=args.insert)
    elif args.command == "xref":
        result = format_document(args.file, "xref")
    elif args.command == "format":
        result = format_document(args.file, "format", preset=args.preset)
    elif args.command == "all":
        result = format_document(args.file, "all", preset=args.preset,
                                 numbering_style=args.numbering_style,
                                 output_path=args.output or args.file)
    elif args.command == "preview":
        result = format_document(args.file, "preview", preset=args.preset)
    else:
        result = {"success": False, "error": "未知命令"}

    print(json.dumps(result, ensure_ascii=False, default=str, indent=2))
